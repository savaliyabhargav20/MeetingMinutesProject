import streamlit as st
from docx import Document
from io import BytesIO

# 1. Define the missing function BEFORE calling it
def create_docx(text):
    doc = Document()
    doc.add_heading('Meeting Summary', 0)
    doc.add_paragraph(text)
    
    # Save to a buffer so we can download it in Streamlit
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

st.title("Meeting Minutes Project")

# ... (Your logic for generating the summary goes here) ...
summary = "This is a sample summary of the meeting." 

# 2. Now the script knows what 'create_docx' is
if st.button("Generate Word Document"):
    docx_file = create_docx(summary)
    
    st.download_button(
        label="Download Report",
        data=docx_file,
        file_name="meeting_summary.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

st.set_page_config(page_title="AI Meeting Minutes", page_icon="📝")
st.title("📝 Automatic Meeting Minutes")

uploaded_file = st.file_uploader("Upload Zoom Recording", type=["m4a", "mp3", "wav"])

if uploaded_file:
    if st.button("Generate Minutes"):
        # Create a temp file to process
        with open("temp_audio.m4a", "wb") as f:
            f.write(uploaded_file.getbuffer())
        
        with st.spinner("Processing..."):
            text = transcribe_audio("temp_audio.m4a")
            minutes = generate_summary(text)
            
            st.subheader("Results")
            st.write(minutes)
        
        os.remove("temp_audio.m4a")

openai.api_key = st.secrets["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"]
def transcribe_audio(file_path):
    try:
        # If using OpenAI API
        with open(file_path, "rb") as audio_file:
            transcript = openai.Audio.transcribe("whisper-1", audio_file)
        return transcript["text"]
    except Exception as e:
        return f"Error during transcription: {str(e)}"

# --- 2. YOUR APP LOGIC ---
st.title("Meeting Minutes Generator")

uploaded_file = st.file_uploader("Upload meeting audio", type=["m4a", "mp3", "wav"])

if uploaded_file is not None:
    # Save the uploaded file temporarily
    with open("temp_audio.m4a", "wb") as f:
        f.write(uploaded_file.getbuffer())

    st.success("Audio uploaded! Starting transcription...")

    # --- THIS IS LINE 45 WHERE THE ERROR WAS ---
    text = transcribe_audio("temp_audio.m4a")
    
    st.write("### Transcript:")
    st.write(text)


