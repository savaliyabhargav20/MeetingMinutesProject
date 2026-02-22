import whisper
import os
import io
from docx import Document
from dotenv import load_dotenv
from langchain_openai import ChatOpenAI
from langchain_core.prompts import ChatPromptTemplate

load_dotenv()

def transcribe_audio(file_path):
    """Converts audio to text using OpenAI Whisper."""
    try:
        # 'base' is good for speed/accuracy balance
        model = whisper.load_model("base")
        result = model.transcribe(file_path)
        return result["text"]
    except Exception as e:
        raise Exception(f"Transcription failed: {str(e)}")

def generate_summary(transcript_text):
    """Uses a professional prompt to summarize meeting notes."""
    try:
        llm = ChatOpenAI(model="gpt-4o", temperature=0.5)
        
        prompt = ChatPromptTemplate.from_template("""
        You are a professional secretary. Summarize the following meeting transcript into:
        1. **Executive Summary** (brief overview)
        2. **Key Decisions** (bullet points)
        3. **Action Items** (Table format: Task | Owner)
        
        Transcript:
        {transcript}
        """)
        
        chain = prompt | llm
        response = chain.invoke({"transcript": transcript_text})
        return response.content
    except Exception as e:
        raise Exception(f"Summarization failed: {str(e)}")

def create_docx(summary_text):
    """Creates a professional Word document from the summary."""
    doc = Document()
    doc.add_heading('Meeting Minutes', 0)
    doc.add_paragraph(summary_text)
    
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()
